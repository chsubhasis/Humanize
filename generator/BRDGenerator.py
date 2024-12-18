import docx
from langchain_core.prompts import PromptTemplate
from langchain_core.prompts import FewShotPromptTemplate
import os
from typing import List
from mistralai import Mistral
from processor.DocumentProcessor import SAPDocumentProcessor
import traceback

class SAPBRDGenerator:
    def __init__(self, api_key: str, model: str = 'mistral-large-latest', temperature=0.3):
        self.client = Mistral(api_key=api_key)
        self.model = model
        self.temperature = temperature
        # Standard BRD Sections
        self.standard_sections = [
            "1. Executive Summary",
            "2. Project Scope",
            "3. Business Requirements",
            "4. Functional Requirements", 
            "5. Non-Functional Requirements",
            "6. Constraints and Assumptions",
            "7. Stakeholder Requirements",
            "8. High-Level Solution Architecture",
            "9. Risk Analysis",
            "10. Acceptance Criteria"
        ]
        
        # Conversation context
        self.current_assessment = None
        self.current_brd = None
        
        # Few-shot examples
        self.few_shot_examples = []

    def load_few_shot_examples(self, assessment_paths: List[str], brd_paths: List[str]) -> None:
        """
        Load few-shot learning examples from assessment and BRD document pairs.
        
        Args:
            assessment_paths (List[str]): Paths to assessment documents
            brd_paths (List[str]): Corresponding paths to BRD documents
        """
        self.few_shot_examples = []
        
        for assess_path, brd_path in zip(assessment_paths, brd_paths):
            try:
                # Extract text from assessment and BRD
                assessment_text = SAPDocumentProcessor.extract_text(assess_path)
                brd_text = SAPDocumentProcessor.extract_text(brd_path)
                
                self.few_shot_examples.append({
                    'input': assessment_text,
                    'output': brd_text
                })
            except Exception as e:
                error_message = f"Error: {str(e)}\n\nStack Trace:\n{traceback.format_exc()}"
                return error_message, None

    def create_few_shot_prompt_template(self) -> FewShotPromptTemplate:
        """
        Create a sophisticated few-shot prompt template for BRD generation.
        
        Returns:
            FewShotPromptTemplate: Configured prompt template
        """
        # Example prompt template
        example_prompt = PromptTemplate(
            input_variables=["input", "output"],
            template="""
Sample Assessment Report:
{input}

Corresponding Sample Business Requirements Document:
{output}
"""
        )
        
        # Prompt template for new BRD generation
        main_prompt = PromptTemplate(
            input_variables=["assessment_report"],
            template="""
You are an expert SAP Business Requirements Document (BRD) generator.

Carefully analyze the following assessment report and generate a comprehensive BRD.

Key Guidelines:
- Use clear, concise, and professional language
- Directly reference the uploaded assessment report
- Ensure each standard section is thoroughly addressed
- Maintain the structure of previous successful BRDs
- Focus on specific, measurable requirements

Standard Sections to Include:
{sections}

Assessment Report:
{assessment_report}

Generate a comprehensive Business Requirements Document:
"""
        )
        
        # Create Few-Shot Prompt Template
        few_shot_prompt = FewShotPromptTemplate(
            example_prompt=example_prompt,
            examples=self.few_shot_examples,
            input_variables=["assessment_report"],
            prefix="Use the following examples as a guide for generating the BRD:",
            suffix=main_prompt.template.format(
                sections="\n".join(self.standard_sections),
                assessment_report="{assessment_report}"
            )
        )
        
        return few_shot_prompt

    def generate_brd(self, assessment_text: str) -> str:
        """Generate initial Business Requirements Document with few-shot learning."""
        # Reset context for new document
        self.current_assessment = assessment_text

        # Create few-shot prompt template
        few_shot_template = self.create_few_shot_prompt_template()
        
        # Prepare full prompt
        full_prompt = few_shot_template.format(assessment_report=assessment_text)
        
        # Generate BRD
        response = self.client.chat.complete(
            model=self.model,
            messages=[{"role": "user", "content": full_prompt}],
            temperature=self.temperature
        )
        
        # Store and return BRD content
        self.current_brd = response.choices[0].message.content
        return self.current_brd

    def refine_brd(self, user_feedback: str) -> str:
        """Refine BRD based on user feedback."""
        if not self.current_brd:
            return "No existing BRD to refine."

        # Prepare context with previous interactions
        messages = [
            {
                "role": "system", 
                "content": "Refine the Business Requirements Document based on user feedback."
            },
            {
                "role": "user", 
                "content": f"Original Assessment: {self.current_assessment}"
            },
            {
                "role": "assistant", 
                "content": f"Here is the current version of the Business Requirements Document (BRD). Please update it based on the feedback. {self.current_brd}"
            },
            {
                "role": "user", 
                "content": f"Feedback to incorporate: {user_feedback}"
            }
        ]
        
        # Generate refined BRD
        response = self.client.chat.complete(
            model=self.model,
            messages=messages,
            temperature=self.temperature
        )
        
        # Update context
        self.current_brd = response.choices[0].message.content
        #self.chat_history.append({"user": user_feedback, "assistant": self.current_brd})
        
        return self.current_brd

    def save_brd(self, brd_content: str, filename: str = 'generated_brd.docx') -> str:
        doc = docx.Document()
        doc.add_heading('Generated Business Requirements Document', level=1)
        doc.add_paragraph(brd_content)

        # Ensure directory exists
        os.makedirs('generated_brds', exist_ok=True)
        filepath = os.path.join('generated_brds', filename)
        doc.save(filepath)
        
        return filepath